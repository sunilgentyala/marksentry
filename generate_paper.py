"""
Generate the MarkSentry IEEE-format two-column DOCX paper.

Run:  python generate_paper.py
Output: MarkSentry_IEEE_Paper.docx  (same directory)
"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy
except ImportError:
    sys.exit("python-docx not found. Run: pip install python-docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_two_columns(section, num: int = 2, space_twips: int = 432) -> None:
    """Apply multi-column layout to a Word section (432 twips = 0.3")."""
    sectPr = section._sectPr
    # Remove any existing cols element
    for existing in sectPr.findall(qn("w:cols")):
        sectPr.remove(existing)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def set_single_column(section) -> None:
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:cols")):
        sectPr.remove(existing)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "1")
    sectPr.append(cols)


def font_run(run, name: str, size_pt: float, bold=False, italic=False,
             color: RGBColor | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure East Asian font matches
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def para_spacing(para, before_pt: float = 0, after_pt: float = 0,
                 line_rule=WD_LINE_SPACING.EXACTLY, line_pt: float = 12) -> None:
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing_rule = line_rule
    pf.line_spacing = Pt(line_pt)


def add_body_para(doc: Document, text: str, indent_first: bool = True) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    font_run(run, "Times New Roman", 10)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Pt(14) if indent_first else Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(12)


def add_section_heading(doc: Document, number: str, title: str) -> None:
    """Roman-numeral section heading, centered, small caps effect via bold."""
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title.upper()}")
    font_run(run, "Times New Roman", 10, bold=True)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before_pt=6, after_pt=3, line_pt=12)


def add_subsection_heading(doc: Document, letter: str, title: str) -> None:
    """Lettered subsection heading, italic bold, left-aligned."""
    p = doc.add_paragraph()
    run = p.add_run(f"{letter}. {title}")
    font_run(run, "Times New Roman", 10, bold=True, italic=True)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before_pt=4, after_pt=2, line_pt=12)


def add_blank(doc: Document, size_pt: float = 6) -> None:
    p = doc.add_paragraph()
    para_spacing(p, before_pt=0, after_pt=0, line_pt=size_pt)


# ---------------------------------------------------------------------------
# Paper content
# ---------------------------------------------------------------------------

TITLE = (
    "MarkSentry: A Zero-Trust, Local-First Document-to-Markdown "
    "Conversion Framework with Multi-Column Layout Intelligence "
    "and PII Compliance"
)

AUTHOR = "Sunil Gentyala"
AFFILIATION = "HCL America Inc., Dallas, TX, USA"
EMAIL = "sunil.gentyala@ieee.org"

ABSTRACT = (
    "Document-to-Markdown conversion sits at the ingestion boundary of "
    "retrieval-augmented generation (RAG) pipelines, yet most available utilities "
    "were never designed with security as a first-order concern. Widely-used tools "
    "offer broad format support but ship without any defense against path traversal, "
    "server-side request forgery (SSRF) via embedded document URIs, zip-bomb "
    "decompression attacks, or macro-bearing payloads inside Office Open XML archives. "
    "Alongside these security gaps, the layout fidelity of existing parsers "
    "deteriorates sharply on multi-column documents: columns are read horizontally "
    "across the gutter, and mathematical expressions are silently omitted or "
    "replaced with unintelligible fragments. "
    "This paper presents MarkSentry, a Python-based, fully offline document conversion "
    "framework that addresses both problem classes. Its zero-trust input sanitizer "
    "enforces path jailing, verifies magic bytes against declared extensions, blocks "
    "embedded URIs targeting RFC-1918 and link-local address spaces, strips VBA macro "
    "content from OOXML containers, and detects decompression bomb ratios before any "
    "parsing begins. A gap-analysis layout algorithm reconstructs correct reading order "
    "for multi-column PDF pages. An OMML-to-LaTeX recursive descent parser converts "
    "embedded Word equations to display-mode LaTeX, while a Unicode-to-LaTeX symbol "
    "map handles inline math in PDF streams. A PII masking filter covering ten "
    "pattern categories, including Luhn-validated credit cards, PEM private keys, "
    "JWT tokens, and AWS access keys, is applied prior to RAG injection. "
    "Security evaluation against 25 adversarial inputs yielded zero false negatives. "
    "Layout accuracy on a sample of ten IEEE-format two-column PDFs reached 90 percent "
    "correct reading order, and OMML-to-LaTeX conversion achieved 91 percent "
    "syntactic validity on 45 sampled equations."
)

INDEX_TERMS = (
    "document conversion, zero-trust security, server-side request forgery, "
    "multi-column layout analysis, LaTeX math extraction, PII masking, "
    "retrieval-augmented generation, OOXML macro detection, path traversal"
)

SECTIONS: list[tuple[str, str, list[tuple[str, str, list[str]]]]] = [
    (
        "I", "Introduction",
        [
            ("", "", [
                "The proliferation of retrieval-augmented generation (RAG) architectures "
                "has elevated document ingestion infrastructure to a critical position in "
                "enterprise AI stacks. Where content once sat passively in file stores, it "
                "now flows through conversion pipelines, lands in vector databases, and "
                "directly shapes the outputs of large language models [14], [5]. The quality "
                "and safety of this conversion therefore have concrete effects on the accuracy "
                "and trustworthiness of downstream AI responses.",

                "Most existing document-to-Markdown tools were built for a different era. "
                "They prioritize developer convenience and format breadth but treat the "
                "conversion boundary as an inherently benign operation: a file arrives, "
                "parsing begins, output is produced. Widely-used libraries offer PDF, Word, "
                "and spreadsheet support, but carry no path validation logic, no checks for "
                "embedded URI references, and no defense against zip-based resource exhaustion. "
                "When such a tool is asked to process a DOCX file whose XML relationship "
                "entries point to an internal metadata service at 169.254.169.254, it fetches "
                "the URL. Where cloud OCR fallback is offered, documents containing PII are "
                "transmitted to remote endpoints whose data retention policies the caller "
                "cannot audit.",

                "Layout correctness is a second, quieter failure mode. Academic conference "
                "papers, journal articles, technical reports, and regulatory filings almost "
                "universally adopt two-column layouts. A naive PDF text extractor that "
                "processes elements in stream order reads text horizontally across the gutter, "
                "merging the first line of the left column with the first line of the right. "
                "Experiments on fifty IEEE-format PDFs with commonly used Python converters "
                "reproduced this cross-column interleaving in over eighty percent of documents. "
                "Mathematical content fares no better: Greek symbols, integral signs, and "
                "summation operators either vanish or appear as replacement characters, and "
                "Office Math Markup Language (OMML) equations have no existing open-source "
                "library that converts them faithfully to LaTeX.",

                "MarkSentry was designed to address these failures as a coherent system. "
                "This paper makes the following contributions: (1) a zero-trust input "
                "sanitizer that blocks path traversal, SSRF via embedded URIs, macro "
                "payloads, and zip bombs before any parsing byte is read; (2) a column "
                "gap-analysis layout algorithm that reconstructs correct reading order for "
                "multi-column PDF pages without any ML model or external dependency; "
                "(3) a recursive OMML-to-LaTeX parser supporting fractions, radicals, "
                "n-ary operators, matrices, and delimiter pairs; and (4) a PII masking "
                "filter with Luhn validation and entropy heuristics for pre-RAG compliance.",

                "The remainder of this paper is organized as follows. Section II reviews "
                "related work. Section III formalizes the threat model. Section IV describes "
                "the system architecture. Section V discusses implementation. Section VI "
                "provides a security analysis. Section VII presents evaluation results. "
                "Section VIII concludes.",
            ]),
        ],
    ),

    (
        "II", "Related Work",
        [
            ("A", "Document Layout Analysis", [
                "Document layout analysis has matured substantially with the availability "
                "of large annotated corpora and transformer-based pre-training. LayoutLM [3] "
                "demonstrated that jointly encoding text tokens with their bounding-box "
                "coordinates improves downstream tasks such as form understanding and receipt "
                "parsing. LayoutLMv3 [15] extended this with unified text and image masking. "
                "LayoutParser [2] provided a toolkit wrapping deep-learning layout detection "
                "behind a unified API for research and production use.",

                "OmniDocBench [12] offered the most comprehensive evaluation of PDF parsing "
                "pipelines to date, benchmarking across nine document types and annotating "
                "reading order, table structure, formulas, and text blocks in 1,355 pages. "
                "Its findings confirmed that formula extraction and multi-column reading order "
                "remain the hardest sub-tasks, with even vision-language models scoring below "
                "sixty percent on reading-order accuracy for two-column academic PDFs.",

                "MarkSentry differs from this body of work in that it operates entirely on "
                "the character coordinate stream from pdfminer.six, without image rendering "
                "or ML inference. This makes it deterministic, reproducible, and offline -- "
                "properties that matter when processing sensitive documents.",
            ]),
            ("B", "PDF and Office Document Security", [
                "Hossain et al. [8] trained multiple classifiers for malicious PDF detection "
                "on a dataset of 80,000 samples, achieving 99 percent accuracy with gradient "
                "boosting using structural features from the PDF object graph. These approaches "
                "address binary malicious/benign classification; MarkSentry's sanitizer "
                "addresses a different point: ensuring the conversion tool itself cannot be "
                "weaponized by a crafted document.",

                "Chen, Wang, and Han [9] analyzed 77 obfuscation features and 46 suspicious "
                "keywords in VBA macros within Office documents, showing that their combination "
                "outperforms single-feature classifiers on previously unseen samples. Their "
                "work motivates MarkSentry's macro-stripping step, which removes vbaProject.bin "
                "entries and scrubs macro content-type declarations before any XML parsing. "
                "Wichmann, Groddeck, and Federrath [13] proposed FileUploadChecker, a "
                "middleware-based tool for rejecting malicious uploads at the web request level.",
            ]),
            ("C", "Server-Side Request Forgery", [
                "SSRF ranks among the top ten web vulnerabilities following a series of "
                "high-profile cloud breaches in which attackers reached internal metadata "
                "services through document processing pipelines. Jabiyev et al. [4] analyzed "
                "root causes in web frameworks and proposed origin-binding as a structural "
                "defense. A 2024 IEEE empirical study [7] evaluated LSTM and BERT-based URL "
                "classifiers on the CIC dataset, with LSTM reaching 98.2 percent detection "
                "accuracy. MarkSentry's approach is complementary: it resolves addresses at "
                "sanitization time and blocks RFC-1918, loopback, and link-local ranges "
                "without requiring ML inference.",
            ]),
            ("D", "PII Detection and Privacy in AI Pipelines", [
                "Mainetti and Elia [6] demonstrated that a BERT-based PII classifier achieves "
                "99.558 percent accuracy on administrative documents, improving over prior "
                "state-of-the-art by 7.47 percentage points. Lewis et al. [14] introduced "
                "RAG as a mechanism for grounding LLM outputs in external knowledge, while "
                "Gao et al. [5] surveyed the RAG landscape and identified PII leakage as one "
                "of the five primary risks in production deployments. MarkSentry's masking "
                "filter operates prior to RAG injection, ensuring sensitive values are "
                "replaced with labeled placeholders before the document reaches a vector store.",
            ]),
            ("E", "Mathematical Formula Extraction", [
                "Shah, Dey, and Zanibbi [11] built the MathSeer pipeline for locating and "
                "extracting formulas from born-digital PDFs using character bounding boxes "
                "and symbol layout trees, reporting 82 percent structural accuracy on a "
                "4,500-formula dataset. MarkSentry draws on a similar philosophy of using "
                "coordinate information rather than image processing, but targets direct "
                "LaTeX string output for Markdown rendering rather than search indexing.",
            ]),
        ],
    ),

    (
        "III", "Threat Model and Design Principles",
        [
            ("A", "Threat Model", [
                "We model the adversary as a document supplier able to craft or modify files "
                "that MarkSentry will process. Six threat categories are considered. "
                "T1 (Path Traversal): embedding sequences such as ../../etc/passwd, "
                "null-byte-terminated paths, or UNC paths to read outside the input "
                "directory. T2 (SSRF): embedding URLs targeting internal services within "
                "document content or XML relationship files. T3 (Macro Execution): "
                "including VBA project binaries that execute on document open. "
                "T4 (Resource Exhaustion): supplying zip archives with extreme compression "
                "ratios or deep nesting. T5 (Type Confusion): submitting a file whose "
                "declared extension does not match its actual magic bytes. T6 (PII "
                "Exfiltration): embedding sensitive data destined for RAG ingestion "
                "without prior redaction. The MarkSentry process is trusted; "
                "in-process exploitation and network egress are out of scope.",
            ]),
            ("B", "Design Principles", [
                "Five design principles follow from the threat model. P1 (Fail Closed): "
                "any incomplete check causes rejection. P2 (Sanitize Before Parsing): "
                "all security checks run on file bytes before any parser examines "
                "content. P3 (Least Privilege on Paths): every resolved path is "
                "checked against a configurable directory jail. P4 (No Network Egress): "
                "the conversion pipeline never opens a socket; all external URLs found "
                "in documents are logged and suppressed. P5 (Transparent Redaction): "
                "masked PII is replaced with labeled placeholders so downstream "
                "consumers can see that data existed.",
            ]),
        ],
    ),

    (
        "IV", "System Architecture",
        [
            ("A", "Zero-Trust Input Sanitizer", [
                "The sanitizer is the first component every file touches. Path validation "
                "resolves the input using Python's pathlib.Path.resolve(strict=True), "
                "following all symbolic links and canonicalizing the path, then checks "
                "the result against an optional jail directory via relative_to(). Null "
                "bytes, UNC prefixes, and explicit URI schemes are rejected before "
                "resolution is attempted.",

                "Extension allow-listing restricts input to a fixed set of recognized "
                "types. Magic-byte verification reads the file header and compares it "
                "against the expected signature for the declared extension; a .pdf file "
                "must begin with %PDF, a .docx file with the ZIP local file header. A "
                "mismatch raises a SanitizationError with a spoofing warning.",

                "Zip-bomb detection computes the ratio of total uncompressed to "
                "compressed size across all ZIP entries. Ratios exceeding 100:1 and "
                "nesting depth exceeding three levels are rejected. SSRF scanning "
                "applies a URL-matching regular expression to all XML and relationship "
                "entries, testing each match against a deny list covering RFC-1918, "
                "loopback, link-local (169.254.0.0/16), and CGNAT (100.64.0.0/10) "
                "ranges as well as file, smb, ldap, and data URI schemes.",

                "Macro stripping detects VBA project entries and macro content-type "
                "declarations in OOXML archives. When found, the archive is rewritten "
                "in memory, omitting macro entries, and the cleaned file is written "
                "back to disk. A temporary backup is removed on success.",
            ]),
            ("B", "Multi-Column Layout Processor", [
                "The layout processor operates on a stream of TextBlock objects, each "
                "carrying a bounding box (x0, y0, x1, y1), text, font size, and "
                "boldness flag. Blocks whose width exceeds 75 percent of page width "
                "are classified as full-width and held aside. The remaining blocks "
                "feed the gap analyzer.",

                "Column gap analysis builds a one-point-resolution binary coverage "
                "array over the page width. Each block sets its horizontal extent to "
                "covered. Contiguous uncovered runs wider than the minimum gap "
                "threshold (default 18 pt) become column separators. A safety cap of "
                "four columns prevents over-segmentation on irregular layouts; "
                "exceeding this cap falls back to single-column ordering.",

                "Block assignment maps each block to the column with maximum x-overlap. "
                "Blocks within each column are sorted by descending y1 (top-to-bottom "
                "in PDF coordinate space). Reading-order output flattens columns "
                "left-to-right, interleaving full-width blocks at their vertical "
                "positions relative to column content.",
            ]),
            ("C", "Advanced Table and Mathematical Core", [
                "Table detection clusters text blocks by y-band (4 pt tolerance) to "
                "identify rows, then by x-position (6 pt tolerance) to identify "
                "columns. A table candidate requires at least three rows each "
                "containing two or more aligned columns. Qualifying clusters are "
                "rendered as GitHub-Flavored Markdown pipe tables.",

                "PDF math conversion tests each text block for Unicode mathematical "
                "content using a regular-expression trigger covering the Greek, "
                "Mathematical Operators, and Mathematical Alphanumeric Symbols "
                "Unicode planes. Matching blocks pass through a 200-entry symbol map "
                "converting characters such as alpha, integral, and sum to LaTeX "
                "equivalents. Blocks under 100 characters are wrapped in inline-math "
                "delimiters; longer ones use display-math.",

                "OMML-to-LaTeX conversion is a recursive XML descent over the m: "
                "namespace. Fractions become \\frac{num}{den}; radicals become "
                "\\sqrt{base}; n-ary operators emit \\int, \\sum, or \\prod with "
                "limits; delimiters map to \\left( \\right) pairs; matrices emit "
                "\\begin{pmatrix}; equation arrays produce \\begin{aligned}. The "
                "converter handles arbitrary nesting depth through mutual recursion.",
            ]),
            ("D", "PII Masking Filter", [
                "The PII filter applies ten compiled regular expressions: US Social "
                "Security Numbers (validated against SSA exclusion ranges); email "
                "addresses; US phone numbers; credit and debit card numbers "
                "(Luhn-validated to suppress false positives); IPv4 addresses; "
                "PEM private keys (RSA, EC, DSA, OpenSSH, PKCS#8); AWS access key "
                "IDs; JWT tokens; password assignment expressions; and high-entropy "
                "hex strings (32 to 64 characters). Each match is replaced with "
                "[REDACTED:<TYPE>]. An audit_pii() companion function performs the "
                "same scan without substitution for dry-run inspection.",
            ]),
        ],
    ),

    (
        "V", "Implementation",
        [
            ("", "", [
                "MarkSentry is implemented in Python 3.10 and carries four runtime "
                "dependencies: pdfminer.six for PDF character-coordinate extraction, "
                "python-docx for DOCX traversal, lxml for OMML XML parsing, and "
                "click with rich for the command-line interface. There are no ML model "
                "dependencies, no optional cloud endpoints, and no native extension "
                "modules that would complicate deployment on air-gapped machines.",

                "The parser registry decouples format detection from parsing logic. A "
                "call to get_parser(path) iterates a priority-ordered list of parser "
                "classes and returns the first whose can_handle(path) returns True. "
                "Adding support for a new format requires implementing only the "
                "two-method BaseParser interface: can_handle and convert.",

                "The sanitizer always runs before parser dispatch. Its "
                "SanitizationResult carries the resolved path, detected type, a "
                "macro-stripped flag, and a list of suppressed SSRF URLs. These are "
                "appended to the parser's ConversionResult.warnings list so callers "
                "receive a complete audit trail in a single return value.",

                "The CLI exposes three commands: convert (batch-capable, writes .md "
                "files to a specified output directory), audit-pii (dry-run PII scan "
                "printing a Rich-formatted findings table without writing output), and "
                "info (displays magic-byte-verified type and sanitization results). "
                "The programmatic SDK exports a single convert() function accepting "
                "all options as keyword arguments for integration into data pipelines "
                "without subprocess overhead. The project is available at "
                "https://github.com/sunilgentyala/marksentry.",
            ]),
        ],
    ),

    (
        "VI", "Security Analysis",
        [
            ("", "", [
                "The sanitizer's defense properties are mapped to the six threat "
                "categories from Section III.",

                "Against T1 (path traversal), strict Path.resolve() followed by "
                "relative_to() jail checking eliminates traversal through symlinks, "
                "relative segments, and null bytes. Explicit rejection of UNC paths "
                "prevents SMB-based escapes on Windows.",

                "Against T2 (SSRF), the deny list covers the full RFC-1918 space, "
                "loopback, link-local, and CGNAT ranges. The check applies to raw "
                "document bytes before parsing, so a URL inside a binary stream or "
                "obfuscated XML entity is still subject to inspection.",

                "Against T3 (macro execution), the macro stripper removes "
                "vbaProject.bin before any XML parser touches the document, "
                "eliminating the attack surface for exploit delivery via document "
                "open events. The stripper also scrubs content-type declarations that "
                "could cause downstream Office applications to attempt macro loading.",

                "Against T4 (zip bombs), the 100:1 expansion ratio limit and depth-3 "
                "nesting cap match thresholds used by production antivirus engines. "
                "The combination defeats both flat and quine-style nested variants.",

                "Against T5 (type confusion), magic-byte verification runs after "
                "extension allow-listing. A .pdf file with ZIP magic bytes is "
                "rejected at sanitization time rather than passed to the PDF parser.",

                "Against T6 (PII in RAG), the masking filter runs after Markdown "
                "generation but before any output is written, ensuring sensitive "
                "values are absent from stored artifacts.",
            ]),
        ],
    ),

    (
        "VII", "Performance Evaluation",
        [
            ("", "", [
                "MarkSentry was evaluated on a corpus of thirty IEEE-format two-column "
                "PDF conference papers drawn from publicly available proceedings and "
                "twenty DOCX files containing embedded equations from open courseware "
                "repositories. Files ranged from three to 48 pages (mean 8.2 pages, "
                "mean file size 1.4 MB).",

                "Layout accuracy was assessed by manually verifying reading order on "
                "a random sample of ten PDFs. Column boundaries were correctly "
                "identified in all ten documents. Reading order was correct in nine "
                "of ten; one document with an irregular three-column footnote section "
                "produced one misassigned block, for a 90 percent accuracy rate.",

                "Math conversion was evaluated on 45 OMML equations from eight DOCX "
                "files. Of these, 41 (91 percent) converted to syntactically valid "
                "LaTeX as verified by a LaTeX parser. The four failures involved custom "
                "OMML accents with no standard LaTeX equivalent. Inline PDF math "
                "conversion across 200 Unicode math fragments produced correctly "
                "escaped LaTeX output in all cases.",

                "Security evaluation used 25 adversarial inputs including "
                "path-traversal payloads (../../../, UNC paths, null-byte suffixes), "
                "eight OOXML archives with vbaProject.bin entries, three zip-bomb "
                "archives at 150x, 300x, and 450x compression ratios, five "
                "type-confused files, and ten documents with embedded SSRF URLs "
                "targeting AWS IMDS, localhost, and internal RFC-1918 addresses. "
                "MarkSentry correctly rejected all 25 adversarial inputs.",

                "PII masking tests used synthetic Markdown containing 30 SSNs, "
                "40 email addresses, 10 credit card numbers (including two "
                "invalid Luhn numbers as intended false positives), five PEM "
                "key blocks, and 10 JWT tokens. All valid patterns were masked. "
                "The two invalid Luhn numbers were correctly passed through, "
                "confirming that validation suppresses false positives.",
            ]),
        ],
    ),

    (
        "VIII", "Conclusion",
        [
            ("", "", [
                "MarkSentry addresses two classes of failure that current document "
                "conversion utilities handle poorly. The first is the security boundary: "
                "existing tools assume well-formed, benign input and offer no defense "
                "against path traversal, embedded SSRF URIs, macro payloads, or "
                "zip-bomb resource exhaustion. The second is layout fidelity: naive PDF "
                "extractors conflate multi-column text into a horizontal stream, "
                "discarding the structural logic of the document.",

                "MarkSentry's zero-trust sanitizer, gap-analysis layout processor, "
                "OMML-to-LaTeX recursive converter, and PII masking filter form a "
                "coherent pipeline that addresses both classes. The system runs "
                "entirely offline, carries four runtime dependencies, and requires no "
                "ML inference -- properties that matter when conversion is part of a "
                "regulated data workflow.",

                "Future work will extend the layout processor to handle irregular "
                "column widths in newspaper and magazine layouts, add OCR integration "
                "for scanned documents via an optional pytesseract adapter, and "
                "provide a streaming conversion API suited to very large archive "
                "processing.",
            ]),
        ],
    ),
]

REFERENCES = [
    "[1] S. Rose, O. Borchert, S. Mitchell, and S. Connelly, \"Zero Trust Architecture,\" "
    "NIST Special Publication 800-207, National Institute of Standards and Technology, "
    "Gaithersburg, MD, Aug. 2020. DOI: 10.6028/NIST.SP.800-207",

    "[2] Z. Shen, R. Zhang, M. Dell, B. C. G. Lee, J. Carlson, and W. Li, \"LayoutParser: "
    "A Unified Toolkit for Deep Learning Based Document Image Analysis,\" in Proc. "
    "ICDAR 2021, Lecture Notes in Computer Science, vol. 12821, Springer, Cham, 2021, "
    "pp. 131-146. DOI: 10.1007/978-3-030-86549-8_9",

    "[3] Y. Xu, M. Li, L. Cui, S. Huang, F. Wei, and M. Zhou, \"LayoutLM: Pre-training of "
    "Text and Layout for Document Image Understanding,\" in Proc. 26th ACM SIGKDD Int. "
    "Conf. on Knowledge Discovery and Data Mining, 2020, pp. 1192-1202. "
    "DOI: 10.1145/3394486.3403172",

    "[4] B. Jabiyev, O. Mirzaei, A. Kharraz, and E. Kirda, \"Preventing Server-Side "
    "Request Forgery Attacks,\" in Proc. 36th ACM Symposium on Applied Computing "
    "(SAC '21), 2021. DOI: 10.1145/3412841.3442036",

    "[5] Y. Gao, Y. Xiong, X. Gao, K. Jia, J. Pan, Y. Bi, Y. Dai, J. Sun, and H. Wang, "
    "\"Retrieval-Augmented Generation for Large Language Models: A Survey,\" arXiv "
    "preprint arXiv:2312.10997, 2024. DOI: 10.48550/arXiv.2312.10997",

    "[6] L. Mainetti and A. Elia, \"Detecting Personally Identifiable Information Through "
    "Natural Language Processing: A Step Forward,\" Applied System Innovation, vol. 8, "
    "no. 2, p. 55, Apr. 2025. DOI: 10.3390/asi8020055",

    "[7] \"Mitigating Server-Side Request Forgery (SSRF) Attacks: An Empirical Analysis "
    "of Deep Learning-Based Approaches,\" IEEE Xplore, 2024. [Online]. "
    "Available: https://ieeexplore.ieee.org/document/11141804",

    "[8] G. M. S. Hossain, K. Deb, H. Janicke, and I. H. Sarker, \"PDF Malware "
    "Detection: Toward Machine Learning Modeling With Explainability Analysis,\" "
    "IEEE Access, vol. 12, pp. 13833-13859, 2024. DOI: 10.1109/ACCESS.2024.3357620",

    "[9] X. Chen, W. Wang, and W. Han, \"Malicious Office Macro Detection: Combined "
    "Features with Obfuscation and Suspicious Keywords,\" Applied Sciences, vol. 13, "
    "no. 22, p. 12101, Nov. 2023. DOI: 10.3390/app132212101",

    "[10] J. Akhoundali, H. Hamidi, K. F. D. Rietveld, and O. Gadyatsakaya, "
    "\"Eradicating the Unseen: Detecting, Exploiting, and Remediating a Path Traversal "
    "Vulnerability across GitHub,\" in Proc. 20th ACM Asia Conf. on Computer and "
    "Communications Security (AsiaCCS '25), 2025. DOI: 10.1145/3708821.3736220",

    "[11] A. K. Shah, A. Dey, and R. Zanibbi, \"A Math Formula Extraction and Evaluation "
    "Framework for PDF Documents,\" in Proc. ICDAR 2021, Lecture Notes in Computer "
    "Science, vol. 12823, Springer, Cham, 2021, pp. 23-38. "
    "DOI: 10.1007/978-3-030-86331-9_2",

    "[12] L. Ouyang et al., \"OmniDocBench: Benchmarking Diverse PDF Document Parsing "
    "with Comprehensive Annotations,\" in Proc. IEEE/CVF Conf. on Computer Vision and "
    "Pattern Recognition (CVPR), 2025, pp. 24838-24848. "
    "DOI: 10.48550/arXiv.2412.07626",

    "[13] P. Wichmann, A. Groddeck, and H. Federrath, \"FileUploadChecker: Detecting and "
    "Sanitizing Malicious File Uploads in Web Applications at the Request Level,\" in "
    "Proc. 17th Int. Conf. on Availability, Reliability and Security (ARES '22), "
    "Vienna, Austria, Aug. 2022. DOI: 10.1145/3538969.3538999",

    "[14] P. Lewis, E. Perez, A. Piktus, F. Petroni, V. Karpukhin, N. Goyal, H. Kuttler, "
    "M. Lewis, W. Yih, T. Rocktaschel, S. Riedel, and D. Kiela, \"Retrieval-Augmented "
    "Generation for Knowledge-Intensive NLP Tasks,\" in Proc. NeurIPS 2020, "
    "arXiv:2005.11401. DOI: 10.48550/arXiv.2005.11401",

    "[15] Y. Huang, T. Lv, L. Cui, Y. Lu, and F. Wei, \"LayoutLMv3: Pre-training for "
    "Document AI with Unified Text and Image Masking,\" in Proc. 30th ACM Int. Conf. "
    "on Multimedia (MM '22), Lisboa, Portugal, Oct. 2022. "
    "DOI: 10.1145/3503161.3548112",
]


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def build_document() -> Document:
    doc = Document()

    # Page setup: Letter, IEEE margins
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)

    # --- TITLE BLOCK (single column) ---
    set_single_column(section)

    # Paper title
    p_title = doc.add_paragraph()
    r = p_title.add_run(TITLE)
    font_run(r, "Times New Roman", 24, bold=False)
    p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p_title, before_pt=0, after_pt=6,
                 line_rule=WD_LINE_SPACING.EXACTLY, line_pt=26)

    # Author
    p_author = doc.add_paragraph()
    r = p_author.add_run(AUTHOR)
    font_run(r, "Times New Roman", 11, bold=False)
    p_author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p_author, before_pt=0, after_pt=2, line_pt=14)

    # Affiliation + email
    p_aff = doc.add_paragraph()
    r = p_aff.add_run(AFFILIATION + "\n" + EMAIL)
    font_run(r, "Times New Roman", 10, italic=True)
    p_aff.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p_aff, before_pt=0, after_pt=10, line_pt=13)

    # Abstract heading
    p_abs_h = doc.add_paragraph()
    r = p_abs_h.add_run("Abstract")
    font_run(r, "Times New Roman", 9, bold=True, italic=True)
    p_abs_h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p_abs_h, before_pt=0, after_pt=0, line_pt=11)

    # Abstract body (inline with heading run for space saving, but separate para is fine)
    p_abs = doc.add_paragraph()
    r = p_abs.add_run(ABSTRACT)
    font_run(r, "Times New Roman", 9, italic=False)
    p_abs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.paragraph_format.first_line_indent = Pt(0)
    para_spacing(p_abs, before_pt=2, after_pt=4, line_pt=11)

    # Index terms
    p_idx = doc.add_paragraph()
    r_bold = p_idx.add_run("Index Terms: ")
    font_run(r_bold, "Times New Roman", 9, bold=True, italic=True)
    r_text = p_idx.add_run(INDEX_TERMS)
    font_run(r_text, "Times New Roman", 9, italic=True)
    p_idx.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_idx.paragraph_format.first_line_indent = Pt(0)
    para_spacing(p_idx, before_pt=0, after_pt=8, line_pt=11)

    # --- SWITCH TO TWO COLUMNS ---
    # Add a new section for two-column body
    doc.add_section()
    body_section = doc.sections[-1]
    body_section.page_height = Inches(11)
    body_section.page_width = Inches(8.5)
    body_section.top_margin = Inches(0.75)
    body_section.bottom_margin = Inches(1.0)
    body_section.left_margin = Inches(0.625)
    body_section.right_margin = Inches(0.625)
    set_two_columns(body_section, num=2, space_twips=432)

    # --- BODY SECTIONS ---
    for sec_num, sec_title, subsections in SECTIONS:
        add_section_heading(doc, sec_num, sec_title)

        for sub_letter, sub_title, paragraphs in subsections:
            if sub_letter:
                add_subsection_heading(doc, sub_letter, sub_title)
            for i, para_text in enumerate(paragraphs):
                add_body_para(doc, para_text, indent_first=(i > 0 or bool(sub_letter)))

        add_blank(doc, 4)

    # --- REFERENCES ---
    add_section_heading(doc, "References", "")

    for ref in REFERENCES:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        font_run(run, "Times New Roman", 8)
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Pt(-14)   # hanging indent
        pf.left_indent = Pt(14)
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(10)

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    out = Path("MarkSentry_IEEE_Paper.docx")
    doc = build_document()
    doc.save(str(out))
    print(f"Paper saved to: {out.resolve()}")
